#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Document Converter

This script processes various file types and generates a consolidated Word report
with embedded content previews, table of contents, and internal navigation.

Author: Original author
Version: 1.0
"""

import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image
import openpyxl
from io import BytesIO
import textwrap
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fnmatch
import win32com.client  # For updating TOC
import docx.opc.constants  # For hyperlink functionality
import pythoncom


def process_pdf(file_path, doc):
    """
    Process a PDF file and add its first page to the Word document.
    
    Args:
        file_path (str): Path to the PDF file
        doc (Document): Word document object
        
    Returns:
        None
    """
    try:
        pdf = fitz.open(file_path)
        page = pdf[0]
        zoom = 2  # zoom factor for better quality
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        img_path = f"temp_{os.path.basename(file_path)}.png"
        pix.save(img_path)
        # Add to Word doc
        doc.add_picture(img_path, width=Inches(6))
        os.remove(img_path)
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")


def process_word(file_path, doc):
    """
    Process a Word document and add its content to the report.
    
    First attempts to convert to PDF and process as image,
    falls back to direct text extraction if conversion fails.
    
    Args:
        file_path (str): Path to the Word document
        doc (Document): Word document object
        
    Returns:
        None
    """
    try:
        # Convert Word to PDF first
        # Initialize COM in this thread
        pythoncom.CoInitialize()
        
        # Create new Word instance without affecting others
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        
        # Convert to absolute path
        abs_path = os.path.abspath(file_path)
        pdf_path = abs_path.rsplit('.', 1)[0] + '.pdf'
        
        # Open and convert
        try:
            wb = word.Documents.Open(abs_path)
            wb.SaveAs(pdf_path, FileFormat=17)  # 17 represents PDF format
            wb.Close()
            
            # Release Word application instead of quitting
            word = None
            pythoncom.CoUninitialize()
            
            # Now process the PDF using existing PDF processing function
            process_pdf(pdf_path, doc)
            
            # Clean up the temporary PDF
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
                
        except Exception as word_error:
            print(f"Error converting Word to PDF: {str(word_error)}")
            
            # Release Word before falling back
            if word:
                word = None
            pythoncom.CoUninitialize()
            
            # Fallback to original Word processing method
            src_doc = Document(file_path)
            
            # Add first page content
            for paragraph in src_doc.paragraphs:
                doc.add_paragraph(paragraph.text)
                
            # Add first page tables
            for table in src_doc.tables:
                new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns))
                new_table.style = 'Table Grid'
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text
                doc.add_paragraph()  # Add spacing after table
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")


def process_excel(file_path, doc):
    """
    Process an Excel file and add its content to the report.
    
    First attempts to convert to PDF, falls back to direct data
    extraction using openpyxl if conversion fails.
    
    Args:
        file_path (str): Path to the Excel file
        doc (Document): Word document object
        
    Returns:
        None
    """
    try:
        # Convert Excel to PDF first
        # Initialize COM in this thread
        pythoncom.CoInitialize()
        
        # Create new Excel instance without affecting others
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        
        # Convert to absolute path
        abs_path = os.path.abspath(file_path)
        pdf_path = abs_path.rsplit('.', 1)[0] + '.pdf'
        
        try:
            # Open and convert
            wb = excel.Workbooks.Open(abs_path)
            wb.ExportAsFixedFormat(0, pdf_path)  # 0 represents PDF format
            wb.Close(False)
            
            # Release Excel application instead of quitting
            excel = None
            pythoncom.CoUninitialize()
            
            # Now process the PDF using existing PDF processing function
            process_pdf(pdf_path, doc)
            
            # Clean up the temporary PDF
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
                
        except Exception as excel_error:
            print(f"Error converting Excel to PDF: {str(excel_error)}")
            
            # Release Excel before falling back
            if excel:
                excel = None
            pythoncom.CoUninitialize()
            
            # Fallback to openpyxl method
            try:
                wb = openpyxl.load_workbook(file_path, data_only=True)
                sheet = wb.active
                
                # Create a table with the first 10 rows and 10 columns of data
                table = doc.add_table(rows=1, cols=0)
                table.style = 'Table Grid'
                
                # Determine max columns to display (up to 10)
                max_col = min(10, sheet.max_column)
                max_row = min(11, sheet.max_row)  # Header + 10 rows
                
                # Add header row
                header_cells = table.rows[0].cells
                for col in range(1, max_col + 1):
                    if len(header_cells) < col:
                        header_cells = table.add_column().cells
                    cell_value = sheet.cell(row=1, column=col).value
                    if cell_value is not None:
                        header_cells[col-1].text = str(cell_value)
                
                # Add data rows
                for row in range(2, max_row + 1):
                    cells = table.add_row().cells
                    for col in range(1, max_col + 1):
                        cell_value = sheet.cell(row=row, column=col).value
                        if cell_value is not None:
                            cells[col-1].text = str(cell_value)
                
                doc.add_paragraph("Excel data extracted directly using openpyxl.")
                
            except Exception as openpyxl_error:
                print(f"Error processing Excel file: {str(openpyxl_error)}")
                doc.add_paragraph(f"Failed to process Excel file: {str(openpyxl_error)}")
                
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")
        doc.add_paragraph(f"Error processing {file_path}: {str(e)}")


def process_image(file_path, doc):
    """
    Process an image file and add it to the document.
    
    Uses PIL to handle various image formats and ensure proper
    embedding in the Word document.
    
    Args:
        file_path (str): Path to the image file
        doc (Document): Word document object
        
    Returns:
        None
    """
    try:
        # Load the image with PIL first to ensure proper handling
        img = Image.open(file_path)
        
        # Save to BytesIO to avoid file system operations
        img_byte_arr = BytesIO()
        img.save(img_byte_arr, format=img.format)
        img_byte_arr.seek(0)
        
        # Add to document with explicit float width
        doc.add_picture(img_byte_arr, width=Inches(6.0))
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")


def process_text(file_path, doc):
    """
    Process a text file and add its content to the document.
    
    Handles various encodings and special formats like CSV.
    Attempts multiple encodings to handle Chinese text properly.
    
    Args:
        file_path (str): Path to the text file
        doc (Document): Word document object
        
    Returns:
        None
    """
    try:
        # Check if this is a CSV file based on extension
        is_csv = file_path.lower().endswith('.csv')
        
        # Try different encodings for Chinese files
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    if is_csv:
                        # For CSV files, create a table representation
                        import csv
                        reader = csv.reader(f)
                        
                        # Get the first rows (up to 10)
                        rows = []
                        for i, row in enumerate(reader):
                            if i < 10:  # Limit to first 10 rows
                                rows.append(row)
                            else:
                                break
                        
                        if rows:
                            # Create a table
                            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                            table.style = 'Table Grid'
                            
                            # Fill the table with data
                            for i, row in enumerate(rows):
                                for j, cell_value in enumerate(row):
                                    if j < len(table.rows[i].cells):  # Ensure we don't exceed column count
                                        table.rows[i].cells[j].text = cell_value
                            
                            doc.add_paragraph(f"CSV data presented as table (first 10 rows). Detected encoding: {encoding}")
                        else:
                            doc.add_paragraph("CSV file appears to be empty.")
                    else:
                        # For regular text files
                        content = ''.join(f.readlines()[:10])  # First 10 lines
                        doc.add_paragraph(content)
                        doc.add_paragraph(f"Detected encoding: {encoding}")
                    
                    # If we get here without an exception, we found the right encoding
                    break
            except UnicodeDecodeError:
                # If this encoding doesn't work, try the next one
                if encoding == encodings[-1]:  # If this was the last encoding in our list
                    doc.add_paragraph(f"Could not decode file with any of the attempted encodings.")
                    print(f"Error processing {file_path}: Could not decode with any of the attempted encodings")
                continue
            except Exception as e:
                # Handle other exceptions
                print(f"Error processing {file_path} with encoding {encoding}: {str(e)}")
                doc.add_paragraph(f"Error processing file: {str(e)}")
                break
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")
        doc.add_paragraph(f"Error processing {file_path}: {str(e)}")


def add_bookmark(paragraph, bookmark_name):
    """
    Add a bookmark to a paragraph.
    
    Args:
        paragraph (Paragraph): Paragraph object to add bookmark to
        bookmark_name (str): Name of the bookmark
        
    Returns:
        str: Name of the bookmark that was added
    """
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark_name)
    tag.append(start)
    
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    tag.append(end)
    
    return bookmark_name


def get_valid_bookmark_name(file_path):
    """
    Convert a file path to a valid bookmark name.
    
    Removes invalid characters and ensures uniqueness with prefixing.
    
    Args:
        file_path (str): File path to convert to bookmark name
        
    Returns:
        str: Valid bookmark name
    """
    # Remove invalid characters and replace with underscores
    name = os.path.basename(file_path)
    name = name.replace('.', '_').replace(' ', '_')
    # Add a prefix to ensure uniqueness
    return f"bm_{name}"


def add_internal_hyperlink(paragraph, anchor_text, bookmark_name, tooltip=None):
    """
    Add an internal hyperlink to a paragraph.
    
    Creates a clickable link to navigate within the document.
    
    Args:
        paragraph (Paragraph): Paragraph object to add hyperlink to
        anchor_text (str): Text to display as the hyperlink
        bookmark_name (str): Name of the bookmark to link to
        tooltip (str, optional): Tooltip text for the hyperlink
        
    Returns:
        None
    """
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    
    if tooltip:
        hyperlink.set(qn('w:tooltip'), tooltip)
    
    new_run = OxmlElement('w:r')
    hyperlink.append(new_run)
    
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    t = OxmlElement('w:t')
    t.text = anchor_text
    new_run.append(t)
    
    paragraph._p.append(hyperlink)


def add_hyperlink(paragraph, text, url):
    """
    Add an external hyperlink to a paragraph.
    
    Creates a clickable link to an external resource.
    
    Args:
        paragraph (Paragraph): Paragraph object to add hyperlink to
        text (str): Text to display as the hyperlink
        url (str): URL to link to
        
    Returns:
        None
    """
    # Add a hyperlink to a paragraph
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    hyperlink.append(new_run)
    
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    paragraph._p.append(hyperlink)


def get_path_context(file_path, levels=3):
    """
    Extract the last N levels of a file path for better context.
    
    Args:
        file_path (str): Full file path
        levels (int, optional): Number of directory levels to include. Defaults to 3.
        
    Returns:
        str: String representing the path context with N levels
    """
    parts = []
    path = os.path.dirname(file_path)
    
    for _ in range(levels):
        head, tail = os.path.split(path)
        if tail:
            parts.insert(0, tail)
        path = head
        if not path or path == os.path.sep:
            break
    
    if parts:
        return os.path.join(*parts)
    return os.path.dirname(file_path)


def generate_report(input_dir, output_file):
    """
    Generate a comprehensive report with previews of all files in the input directory.
    
    Args:
        input_dir (str): Input directory containing files to process
        output_file (str): Path to save the generated report
        
    Returns:
        None
    """
    # Convert output_file to absolute path
    output_file_abs = os.path.abspath(output_file)
    
    doc = Document()
    doc.add_heading('Document Screenshot Report', 0)
    
    # Add table of contents heading with improved formatting
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.style.font.size = Pt(16)
    toc_heading.style.font.bold = True
    
    # Create paragraph for TOC field with better formatting
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # Add TOC field with improved formatting
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    # Enhanced TOC instruction text for better formatting
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u \\w'
    
    # Add separator
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    # Add end field
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    # Assemble TOC field
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    
    # Add spacing after TOC
    doc.add_paragraph()
    doc.add_page_break()
    
    # Dictionary to store file paths and their bookmark names
    file_bookmarks = {}
    
    # Process files and create bookmarks
    for root, _, files in os.walk(input_dir):
        for file in files:
            file_path = os.path.join(root, file)
            ext = file.lower()
            bookmark_name = get_valid_bookmark_name(file_path)
            file_bookmarks[file_path] = bookmark_name
            
            # Get path context (last 3 levels of directory)
            path_context = get_path_context(file_path)
            
            if ext.endswith('.pdf'):
                # Include path context in the heading
                heading_title = f"{os.path.basename(file_path)} [{path_context}]"
                heading_para = doc.add_heading(heading_title, level=2)
                add_bookmark(heading_para, bookmark_name)
                doc.add_paragraph(f"Location: {os.path.dirname(file_path)}")
                process_pdf(file_path, doc)
            elif ext.endswith(('.doc', '.docx')):
                heading_title = f"{os.path.basename(file_path)} [{path_context}]"
                heading_para = doc.add_heading(heading_title, level=2)
                add_bookmark(heading_para, bookmark_name)
                doc.add_paragraph(f"Location: {os.path.dirname(file_path)}")
                process_word(file_path, doc)
            elif ext.endswith(('.xls', '.xlsx')):
                heading_title = f"{os.path.basename(file_path)} [{path_context}]"
                heading_para = doc.add_heading(heading_title, level=2)
                add_bookmark(heading_para, bookmark_name)
                doc.add_paragraph(f"Location: {os.path.dirname(file_path)}")
                process_excel(file_path, doc)
            elif ext.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                heading_title = f"{os.path.basename(file_path)} [{path_context}]"
                heading_para = doc.add_heading(heading_title, level=2)
                add_bookmark(heading_para, bookmark_name)
                doc.add_paragraph(f"Location: {os.path.dirname(file_path)}")
                process_image(file_path, doc)
            elif ext.endswith(('.txt', '.log', '.md', '.csv')):
                heading_title = f"{os.path.basename(file_path)} [{path_context}]"
                heading_para = doc.add_heading(heading_title, level=2)
                add_bookmark(heading_para, bookmark_name)
                doc.add_paragraph(f"Location: {os.path.dirname(file_path)}")
                process_text(file_path, doc)
    
    # Add the file entries with internal links to the document
    doc.add_heading('File Index', level=1)
    
    for file_path, bookmark_name in file_bookmarks.items():
        filename = os.path.basename(file_path)
        file_dir = os.path.dirname(file_path)
        path_context = get_path_context(file_path)
        
        # Format the entry with file name and path context with internal link
        entry_para = doc.add_paragraph()
        add_internal_hyperlink(entry_para, filename, bookmark_name, f"Go to {filename}")
        entry_para.add_run(f" ({os.path.splitext(filename)[1]}) - {path_context}")
        
        # Add the full directory path on a new line with indentation
        entry_para.add_run("\n    Location: " + file_dir)
    
    try:
        doc.save(output_file_abs)
        print(f"Report generated at: {output_file}")
        
        # Update TOC using Word COM interface
        pythoncom.CoInitialize()
        
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        
        # Use absolute path when opening with Word
        word_doc = word.Documents.Open(output_file_abs)
        word_doc.Fields.Update()
        word_doc.Save()
        word_doc.Close()
        
        # Release the Word application instead of quitting
        word = None
        pythoncom.CoUninitialize()
    except PermissionError:
        alt_output = os.path.join(os.getcwd(), f"report_{os.path.basename(output_file)}")
        alt_output_abs = os.path.abspath(alt_output)
        try:
            doc.save(alt_output_abs)
            print(f"Could not save to {output_file} due to permission denied.")
            print(f"Report saved to alternative location: {alt_output}")
            
            # Update TOC using Word COM interface
            pythoncom.CoInitialize()
            
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word_doc = word.Documents.Open(alt_output_abs)
            word_doc.Fields.Update()
            word_doc.Save()
            word_doc.Close()
            
            # Release the Word application instead of quitting
            word = None
            pythoncom.CoUninitialize()
        except Exception as e:
            print(f"Failed to save report to both original and alternative locations.")
            print(f"Please ensure you have write permissions or try a different output location.")
            raise
    except Exception as e:
        print(f"Error saving report: {str(e)}")
        raise


if __name__ == "__main__":
    import argparse
    
    # Set up command line argument parsing
    parser = argparse.ArgumentParser(
        description="Generate a comprehensive report from various file types in a directory",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    
    parser.add_argument("-i", "--input", required=True, 
                      help="Input directory (use quotes for paths with spaces)")
    parser.add_argument("-o", "--output", default="report.docx", 
                      help="Output file name")
    
    args = parser.parse_args()
    
    # Check that the input directory exists
    if not os.path.isdir(args.input):
        print(f"Error: Input directory '{args.input}' does not exist.")
        exit(1)
    
    # Generate the report
    print(f"Processing files in '{args.input}'...")
    generate_report(args.input, args.output)
    print(f"Report generation complete. Output saved to: {os.path.abspath(args.output)}")